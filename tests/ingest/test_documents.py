"""Optional PDF/DOCX document ingest adapter tests (M10 PR-3).

``kosha.ingest.documents`` imports ``pypdf``/``python-docx`` at module scope, so
the core install and ``kosha.ingest`` must never import it -- that boundary is
verified in a subprocess to get a clean, order-independent import graph rather
than relying on whatever this pytest session happened to import already.
Everything that actually exercises extraction is skipped when the optional
``documents`` extra is not installed, so ``uv run pytest`` still passes in a
core (no-extras) environment; those tests import ``docx``/``pypdf`` and
``kosha.ingest.documents`` inside the test body, never at module top, so
collection itself never requires the extra.
"""

from __future__ import annotations

import importlib.util
import subprocess
import sys
import tomllib
from pathlib import Path

import pytest

from kosha.ingest.guardrails import IngestGuardrailError, IngestPolicy
from kosha.model import SourceKind

_DOCUMENTS_EXTRA_INSTALLED = (
    importlib.util.find_spec("pypdf") is not None
    and importlib.util.find_spec("docx") is not None
)

requires_documents_extra = pytest.mark.skipif(
    not _DOCUMENTS_EXTRA_INSTALLED,
    reason="requires the optional `documents` extra (pypdf, python-docx)",
)


def _minimal_pdf_bytes(text: str) -> bytes:
    """Hand-build a minimal single-page PDF, with or without body text.

    Built from raw PDF syntax (no ``pypdf``/``reportlab`` needed to author it)
    so the ``documents`` extra is only required to *read* these fixtures, kept
    tiny and deterministic for pinning the "no extractable text" and "real
    text extraction" contracts precisely.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 612 792] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
    ]
    stream_body = f"BT /F1 24 Tf 100 700 Td ({text}) Tj ET".encode("latin-1") if text else b""
    objects.append(b"<< /Length %d >>\nstream\n" % len(stream_body) + stream_body + b"\nendstream")

    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode("latin-1") + obj + b"\nendobj\n"
    xref_offset = len(out)
    count = len(objects) + 1
    out += f"xref\n0 {count}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    out += f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode(
        "latin-1"
    )
    return bytes(out)


# --- 1. Core import boundary -------------------------------------------------


def _assert_clean_import(import_statement: str) -> None:
    # A fresh subprocess, not sys.modules introspection in-process: other test
    # modules in this same pytest session may have already imported
    # kosha.ingest.documents (and therefore pypdf/docx), which would make an
    # in-process check pass or fail depending on test order rather than on
    # what `import_statement` alone actually pulls in.
    script = (
        f"{import_statement}\n"
        "import sys\n"
        "assert 'pypdf' not in sys.modules, 'pypdf leaked into sys.modules'\n"
        "assert 'docx' not in sys.modules, 'docx leaked into sys.modules'\n"
        "assert 'kosha.ingest.documents' not in sys.modules, "
        "'kosha.ingest.documents leaked into sys.modules'\n"
    )
    result = subprocess.run(
        [sys.executable, "-c", script],
        capture_output=True,
        text=True,
        timeout=30,
    )
    assert result.returncode == 0, result.stderr


def test_importing_kosha_does_not_import_optional_document_dependencies() -> None:
    _assert_clean_import("import kosha")


def test_importing_kosha_ingest_does_not_import_optional_document_dependencies() -> None:
    _assert_clean_import("import kosha.ingest")


def test_pyproject_keeps_document_libraries_behind_the_documents_extra_only() -> None:
    pyproject_path = Path(__file__).resolve().parents[2] / "pyproject.toml"
    with pyproject_path.open("rb") as handle:
        metadata = tomllib.load(handle)

    core_dependencies = metadata["project"]["dependencies"]
    assert not any(
        dep.lower().startswith(("pypdf", "python-docx", "docx")) for dep in core_dependencies
    ), f"pypdf/python-docx must not be unmarked core dependencies, got {core_dependencies}"

    optional_dependencies = metadata["project"]["optional-dependencies"]
    assert "documents" in optional_dependencies
    documents_extra = optional_dependencies["documents"]
    assert any(dep.lower().startswith("pypdf") for dep in documents_extra)
    assert any(dep.lower().startswith("python-docx") for dep in documents_extra)


# --- 2. DOCX extraction: text, provenance, directory ingestion --------------


@requires_documents_extra
def test_ingest_documents_extracts_docx_text_and_sets_document_provenance(
    tmp_path: Path,
) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    docx_path = tmp_path / "policy.docx"
    document = Document()
    document.add_paragraph("Returns are accepted within 30 days.")
    document.add_paragraph("")  # blank paragraphs must be dropped, not joined in
    document.add_paragraph("Shipping is free over $50.")
    document.save(docx_path)

    docs = ingest_documents(docx_path, authority_rank=5)

    assert len(docs) == 1
    doc = docs[0]
    assert doc.text == "Returns are accepted within 30 days.\n\nShipping is free over $50."
    assert doc.source.kind is SourceKind.DOCUMENT
    assert doc.source.source_id == "document:policy.docx"
    assert doc.source.location == "policy.docx"
    assert doc.source.title == "policy"
    assert doc.source.authority_rank == 5


@requires_documents_extra
def test_ingest_documents_over_a_directory_uses_sorted_relative_source_ids(
    tmp_path: Path,
) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    (tmp_path / "a.pdf").write_bytes(_minimal_pdf_bytes("Alpha content"))
    (tmp_path / "ignore.txt").write_text("not a supported document", encoding="utf-8")

    beta = Document()
    beta.add_paragraph("Beta content")
    beta.save(tmp_path / "b.docx")

    nested = tmp_path / "sub"
    nested.mkdir()
    gamma = Document()
    gamma.add_paragraph("Gamma content")
    gamma.save(nested / "c.docx")

    docs = ingest_documents(tmp_path, authority_rank=2)

    ids = [doc.source.source_id for doc in docs]
    assert ids == ["document:a.pdf", "document:b.docx", "document:sub/c.docx"]

    by_id = {doc.source.source_id: doc for doc in docs}
    nested_doc = by_id["document:sub/c.docx"]
    assert nested_doc.source.location == "sub/c.docx"
    assert nested_doc.source.authority_rank == 2
    assert nested_doc.text == "Gamma content"
    assert by_id["document:a.pdf"].text == "Alpha content"


# --- 3. PDF extraction: negative and positive text extraction ---------------


@requires_documents_extra
def test_ingest_documents_fails_loud_on_a_pdf_with_no_extractable_text(
    tmp_path: Path,
) -> None:
    from kosha.ingest.documents import DocumentIngestError, ingest_documents

    blank_pdf = tmp_path / "blank.pdf"
    blank_pdf.write_bytes(_minimal_pdf_bytes(""))

    with pytest.raises(DocumentIngestError, match="no extractable text"):
        ingest_documents(blank_pdf)


@requires_documents_extra
def test_ingest_documents_extracts_real_pdf_text(tmp_path: Path) -> None:
    from kosha.ingest.documents import ingest_documents

    pdf_path = tmp_path / "hello.pdf"
    pdf_path.write_bytes(_minimal_pdf_bytes("Hello Kosha"))

    docs = ingest_documents(pdf_path)

    assert len(docs) == 1
    assert docs[0].text == "Hello Kosha"
    assert docs[0].source.kind is SourceKind.DOCUMENT
    assert docs[0].source.source_id == "document:hello.pdf"


# --- 4. Unsupported extension and missing path fail loudly ------------------


@requires_documents_extra
def test_ingest_documents_rejects_an_unsupported_extension(tmp_path: Path) -> None:
    from kosha.ingest.documents import DocumentIngestError, ingest_documents

    text_file = tmp_path / "notes.txt"
    text_file.write_text("plain text is not a supported document format", encoding="utf-8")

    with pytest.raises(DocumentIngestError, match="unsupported document extension"):
        ingest_documents(text_file)


@requires_documents_extra
def test_ingest_documents_fails_loud_on_a_missing_path(tmp_path: Path) -> None:
    from kosha.ingest.documents import ingest_documents

    with pytest.raises(FileNotFoundError, match="not a document file or directory"):
        ingest_documents(tmp_path / "does-not-exist")


# --- 5. Tight max_bytes fails before extraction is attempted ----------------


@requires_documents_extra
def test_ingest_documents_enforces_max_bytes_before_extracting_docx(tmp_path: Path) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    docx_path = tmp_path / "big.docx"
    document = Document()
    # Deliberately tiny extracted text ("Hi") so only the OOXML container's
    # fixed zip overhead (tens of KB, regardless of content) can trip the
    # cap -- if guard_text's separate text-length check were the one firing
    # instead of the raw byte-read cap, this fixture would not trigger it,
    # keeping the two guards from masking each other.
    document.add_paragraph("Hi")
    document.save(docx_path)
    assert docx_path.stat().st_size > 1000

    # IngestGuardrailError (not a docx-parsing error, not DocumentIngestError)
    # proves the byte cap is enforced on the raw read, before python-docx ever
    # sees a (truncated, unparsable) buffer.
    with pytest.raises(IngestGuardrailError):
        ingest_documents(docx_path, policy=IngestPolicy(max_bytes=100))


@requires_documents_extra
def test_ingest_documents_enforces_max_bytes_before_extracting_pdf(tmp_path: Path) -> None:
    from kosha.ingest.documents import ingest_documents

    pdf_path = tmp_path / "big.pdf"
    pdf_path.write_bytes(_minimal_pdf_bytes("Hello Kosha"))
    assert pdf_path.stat().st_size > 50

    with pytest.raises(IngestGuardrailError):
        ingest_documents(pdf_path, policy=IngestPolicy(max_bytes=50))


# --- 6. Hidden Unicode is stripped from extracted text -----------------------


@requires_documents_extra
def test_ingest_documents_strips_hidden_unicode_from_docx_text(tmp_path: Path) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    docx_path = tmp_path / "poisoned.docx"
    document = Document()
    document.add_paragraph("visible\u200btext")
    document.save(docx_path)

    docs = ingest_documents(docx_path)

    assert docs[0].text == "visibletext"
