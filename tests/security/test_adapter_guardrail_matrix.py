"""Cross-adapter guardrail matrix (M10 PR-4): no adapter bypasses the shared
M6 size/secret/prompt-injection guardrails.

Every per-adapter test file (PR-1 folder, PR-2 workspace exports, PR-3
documents) already proves its own adapter enforces the shared
:class:`~kosha.ingest.guardrails.IngestPolicy` in isolation, each with its own
locally constructed policy and fixture. This module is the regression net
across the *whole* family: one canonical oversized fixture and one canonical
hidden-Unicode fixture, driven through every core adapter with the *same*
policy instance and the *same* poisoned payload, so a future adapter (or a
refactor of an existing one) that forgets to thread the policy through, or
that skips sanitization, fails here even if its own test file is incomplete
or missing.

It also proves the guardrail chain survives end to end: secret-like content
ingested through a folder adapter and a workspace-export adapter reaches the
pipeline's secret scanner and forces the generated change to the BLOCK lane
instead of silently auto-committing -- using the same deterministic local
providers the rest of the suite uses, so nothing here touches the network or
a live SaaS API.
"""

from __future__ import annotations

import importlib.util
import json
from collections.abc import Callable
from datetime import UTC, datetime
from pathlib import Path

import pytest

from kosha.approve import Decision, Lane
from kosha.evidence import hash_evidence_text
from kosha.git_store import GitStore
from kosha.ingest import (
    IngestGuardrailError,
    IngestPolicy,
    ingest_confluence_export,
    ingest_folder,
    ingest_notion_export,
    ingest_slack_export,
)
from kosha.ingest.guardrails import bind_evidence
from kosha.model import RawDoc
from kosha.pipeline import ingest
from kosha.providers import ExtractiveGenerationProvider, LexicalEmbeddingProvider

_ASOF = datetime(2026, 6, 28, tzinfo=UTC)

_AdapterFn = Callable[..., list[RawDoc]]
_BuildRootFn = Callable[[Path], Path]

_DOCUMENTS_EXTRA_INSTALLED = (
    importlib.util.find_spec("pypdf") is not None
    and importlib.util.find_spec("docx") is not None
)
requires_documents_extra = pytest.mark.skipif(
    not _DOCUMENTS_EXTRA_INSTALLED,
    reason="requires the optional `documents` extra (pypdf, python-docx)",
)


def _minimal_pdf_bytes() -> bytes:
    """Hand-build a minimal, syntactically valid, empty single-page PDF.

    No ``pypdf`` needed to author it (only to read it), matching PR-3's
    fixture technique: the shared guardrail runs on raw bytes before any
    parsing is attempted, so validity only matters for the file to plausibly
    look like the format it claims to be.
    """
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{index} 0 obj\n".encode("latin-1") + obj + b"\nendobj\n"
    xref_offset = len(out)
    count = len(objects) + 1
    out += f"xref\n0 {count}\n".encode("latin-1")
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode("latin-1")
    out += f"trailer\n<< /Size {count} /Root 1 0 R >>\nstartxref\n{xref_offset}\n%%EOF".encode(
        "latin-1"
    )
    return bytes(out)


# ---------------------------------------------------------------------------
# 1. A shared, tight IngestPolicy fails loud through every core adapter path.
# ---------------------------------------------------------------------------

_TIGHT_POLICY = IngestPolicy(max_bytes=48)
_OVERSIZED_BODY = "x" * 200  # well over the 48-byte cap once UTF-8 encoded


def _oversized_markdown(tmp_path: Path) -> Path:
    (tmp_path / "big.md").write_text(f"# Big\n\n{_OVERSIZED_BODY}\n", encoding="utf-8")
    return tmp_path


def _oversized_confluence_json(tmp_path: Path) -> Path:
    payload = [{"id": "p1", "title": "Big", "body": _OVERSIZED_BODY}]
    (tmp_path / "pages.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


def _oversized_slack_json(tmp_path: Path) -> Path:
    payload = [{"ts": "1", "user": "agent", "text": _OVERSIZED_BODY}]
    (tmp_path / "general.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


_OVERSIZED_CASES = [
    pytest.param(ingest_folder, _oversized_markdown, id="folder"),
    pytest.param(ingest_confluence_export, _oversized_markdown, id="confluence-markdown"),
    pytest.param(ingest_confluence_export, _oversized_confluence_json, id="confluence-json"),
    pytest.param(ingest_notion_export, _oversized_markdown, id="notion-markdown"),
    pytest.param(ingest_slack_export, _oversized_slack_json, id="slack-json"),
]


@pytest.mark.parametrize("adapter, build_root", _OVERSIZED_CASES)
def test_oversized_fixture_fails_through_every_core_adapter(
    adapter: _AdapterFn, build_root: _BuildRootFn, tmp_path: Path
) -> None:
    root = build_root(tmp_path)
    with pytest.raises(IngestGuardrailError) as exc_info:
        adapter(root, policy=_TIGHT_POLICY)
    # The raised error names *this* policy's own cap, proving the adapter
    # actually threaded the shared instance through to the byte-bounded read
    # rather than silently falling back to some other (e.g. default) policy.
    assert f"max_bytes={_TIGHT_POLICY.max_bytes}" in str(exc_info.value)


@requires_documents_extra
def test_oversized_fixture_fails_through_the_optional_docx_adapter(tmp_path: Path) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    document = Document()
    document.add_paragraph("small")
    docx_path = tmp_path / "big.docx"
    document.save(docx_path)

    with pytest.raises(IngestGuardrailError) as exc_info:
        ingest_documents(docx_path, policy=_TIGHT_POLICY)
    assert f"max_bytes={_TIGHT_POLICY.max_bytes}" in str(exc_info.value)


@requires_documents_extra
def test_oversized_fixture_fails_through_the_optional_pdf_adapter(tmp_path: Path) -> None:
    from kosha.ingest.documents import ingest_documents

    pdf_path = tmp_path / "big.pdf"
    pdf_path.write_bytes(_minimal_pdf_bytes())

    with pytest.raises(IngestGuardrailError) as exc_info:
        ingest_documents(pdf_path, policy=_TIGHT_POLICY)
    assert f"max_bytes={_TIGHT_POLICY.max_bytes}" in str(exc_info.value)


# ---------------------------------------------------------------------------
# 2. Hidden-Unicode prompt-injection obfuscation is stripped by every core
#    adapter, using the exact same poisoned payload across every adapter
#    shape (a plain file, a JSON page body, a JSON message body).
# ---------------------------------------------------------------------------

_HIDDEN_MARKER = "\u200b"  # zero-width space
_POISONED_TEXT = f"visible{_HIDDEN_MARKER}payload{_HIDDEN_MARKER}text"
_SANITIZED_TEXT = "visiblepayloadtext"


def _poisoned_markdown(tmp_path: Path) -> Path:
    (tmp_path / "poisoned.md").write_text(_POISONED_TEXT, encoding="utf-8")
    return tmp_path


def _poisoned_confluence_json(tmp_path: Path) -> Path:
    payload = [{"id": "p1", "title": "Poisoned", "body": _POISONED_TEXT}]
    (tmp_path / "pages.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


def _poisoned_slack_json(tmp_path: Path) -> Path:
    payload = [{"ts": "1", "user": "agent", "text": _POISONED_TEXT}]
    (tmp_path / "general.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


_HIDDEN_UNICODE_CASES = [
    pytest.param(ingest_folder, _poisoned_markdown, id="folder"),
    pytest.param(ingest_confluence_export, _poisoned_markdown, id="confluence-markdown"),
    pytest.param(ingest_confluence_export, _poisoned_confluence_json, id="confluence-json"),
    pytest.param(ingest_notion_export, _poisoned_markdown, id="notion-markdown"),
    pytest.param(ingest_slack_export, _poisoned_slack_json, id="slack-json"),
]


@pytest.mark.parametrize("adapter, build_root", _HIDDEN_UNICODE_CASES)
def test_hidden_unicode_is_stripped_from_every_core_adapters_rawdoc_text(
    adapter: _AdapterFn, build_root: _BuildRootFn, tmp_path: Path
) -> None:
    root = build_root(tmp_path)
    docs = adapter(root)
    assert docs
    text = docs[0].text
    # Both directions matter: a scanner that strips everything (not just the
    # hidden characters) would still pass a bare "marker not in text" check.
    assert _HIDDEN_MARKER not in text
    assert _SANITIZED_TEXT in text


# ---------------------------------------------------------------------------
# 3. Secret-like content, ingested through a folder adapter and a workspace
#    adapter, reaches the pipeline secret scanner and blocks instead of
#    silently auto-committing.
# ---------------------------------------------------------------------------

_SECRET_BODY = (
    "Rotate the deploy key before shipping. "
    "AWS_ACCESS_KEY_ID=AKIAIOSFODNN7EXAMPLE is only for the staging bucket."
)


def _seed_bundle(tmp_path: Path) -> tuple[Path, GitStore, str]:
    bundle = tmp_path / "bundle"
    (bundle / "policies").mkdir(parents=True)
    (bundle / "policies" / "shipping.md").write_text(
        "---\ntype: policy\ntitle: Shipping\n"
        "description: How and when orders ship.\n---\n"
        "Standard shipping ships within 5 business days.\n",
        encoding="utf-8",
    )
    store = GitStore.init(bundle)
    store.commit(["policies/shipping.md"], "chore: seed")
    return bundle, store, store.current_sha("main")


def _secret_folder_docs(tmp_path: Path) -> list[RawDoc]:
    source = tmp_path / "source"
    source.mkdir()
    (source / "runbook.md").write_text(
        f"# Internal Deploy Runbook\n\n{_SECRET_BODY}\n", encoding="utf-8"
    )
    return ingest_folder(source)


def _secret_confluence_docs(tmp_path: Path) -> list[RawDoc]:
    source = tmp_path / "source"
    source.mkdir()
    (source / "runbook.md").write_text(
        f"# Internal Deploy Runbook\n\n{_SECRET_BODY}\n", encoding="utf-8"
    )
    return ingest_confluence_export(source)


@pytest.mark.parametrize(
    "build_docs",
    [
        pytest.param(_secret_folder_docs, id="folder"),
        pytest.param(_secret_confluence_docs, id="confluence-workspace-export"),
    ],
)
def test_secret_like_content_routes_to_block_instead_of_auto_committing(
    build_docs: Callable[[Path], list[RawDoc]], tmp_path: Path
) -> None:
    bundle, store, main_sha = _seed_bundle(tmp_path)
    raw_docs = build_docs(tmp_path)

    result = ingest(
        tmp_path / "unused-source-label",
        bundle,
        asof=_ASOF,
        raw_docs=raw_docs,
        git_store=store,
        branch="ingest/secret-probe",
        embedding_provider=LexicalEmbeddingProvider(),
        generation_provider=ExtractiveGenerationProvider(),
    )

    blocked = result.routing.blocked()
    assert blocked, "expected at least one change routed to BLOCK"
    detectors = {name for route in blocked for name in route.change.secret_detectors}
    assert detectors, "expected a non-empty secret detector on the blocked change"
    assert "aws-access-key-id" in detectors
    assert result.routing.lane is Lane.BLOCK
    assert result.routing.requires_approval is True

    # Default-safe gate: no reader and no --yes, so a blocked plan rejects
    # rather than silently auto-committing the secret-bearing change.
    assert result.decision is Decision.REJECT
    assert result.committed is False
    assert store.current_sha("main") == main_sha
    assert not store.branch_exists("ingest/secret-probe")


# ---------------------------------------------------------------------------
# 4. Every core adapter's RawDocs cross the same evidence boundary: clean text
#    from any adapter shape hashes and binds identically once run through
#    bind_evidence (DEVELOPMENT_PLAN.md M3).
# ---------------------------------------------------------------------------

_CLEAN_BODY = "Ordinary policy prose with no secret-like content at all."


def _clean_markdown(tmp_path: Path) -> Path:
    (tmp_path / "clean.md").write_text(f"# Clean\n\n{_CLEAN_BODY}\n", encoding="utf-8")
    return tmp_path


def _clean_confluence_json(tmp_path: Path) -> Path:
    payload = [{"id": "p1", "title": "Clean", "body": _CLEAN_BODY}]
    (tmp_path / "pages.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


def _clean_slack_json(tmp_path: Path) -> Path:
    payload = [{"ts": "1", "user": "agent", "text": _CLEAN_BODY}]
    (tmp_path / "general.json").write_text(json.dumps(payload), encoding="utf-8")
    return tmp_path


_EVIDENCE_BOUNDARY_CASES = [
    pytest.param(ingest_folder, _clean_markdown, id="folder"),
    pytest.param(ingest_confluence_export, _clean_markdown, id="confluence-markdown"),
    pytest.param(ingest_confluence_export, _clean_confluence_json, id="confluence-json"),
    pytest.param(ingest_notion_export, _clean_markdown, id="notion-markdown"),
    pytest.param(ingest_slack_export, _clean_slack_json, id="slack-json"),
]


@pytest.mark.parametrize("adapter, build_root", _EVIDENCE_BOUNDARY_CASES)
def test_every_core_adapters_docs_bind_to_a_matching_content_digest(
    adapter: _AdapterFn, build_root: _BuildRootFn, tmp_path: Path
) -> None:
    root = build_root(tmp_path)
    docs = adapter(root)
    assert docs
    bound, run = bind_evidence(
        docs,
        run_id="run-parity",
        bundle_identity="a" * 64,
        source_instance_id="parity",
        started_at=_ASOF,
        completed_at=_ASOF,
    )
    assert run is not None
    for raw in bound:
        assert raw.evidence_sha256 == hash_evidence_text(raw.text)
    assert {doc.sha256 for doc in run.run.evidence} == {raw.evidence_sha256 for raw in bound}


@requires_documents_extra
def test_the_optional_document_adapters_docs_bind_to_a_matching_content_digest(
    tmp_path: Path,
) -> None:
    from docx import Document

    from kosha.ingest.documents import ingest_documents

    document = Document()
    document.add_paragraph(_CLEAN_BODY)
    docx_path = tmp_path / "clean.docx"
    document.save(docx_path)

    docs = ingest_documents(docx_path)
    bound, run = bind_evidence(
        docs,
        run_id="run-parity-doc",
        bundle_identity="a" * 64,
        source_instance_id="parity",
        started_at=_ASOF,
        completed_at=_ASOF,
    )
    assert run is not None
    assert bound[0].evidence_sha256 == hash_evidence_text(bound[0].text)
